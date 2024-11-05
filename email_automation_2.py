import datetime
from email.header import decode_header
import smtplib
import imaplib
import schedule
import time
import email
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import sys
import os
from docx import Document
from email.mime.base import MIMEBase
from email import encoders
import os

mailAddress = ""  # enter mail
appPassword = ""  # create and use app password

# configs
gmail_config = {
    "service": "Gmail",
    "smtp_server": "smtp.gmail.com",
    "imap_server": "imap.gmail.com",
    "email_address": mailAddress,
    "password": appPassword,
}

yahoo_config = {
    "service": "Yahoo",
    "smtp_server": "smtp.mail.yahoo.com",
    "imap_server": "imap.mail.yahoo.com",
    "email_address": mailAddress,
    "password": appPassword,
}

outlook_config = {
    "service": "Outlook",
    "smtp_server": "smtp-mail.outlook.com",
    "imap_server": "outlook.office365.com",
    "email_address": mailAddress,
    "password": appPassword,
}


class EmailAutomation: 
    def __init__(self, service, smtp_server, imap_server, email_address, password):
        self.service = service
        self.smtp_server = smtp_server
        self.imap_server = imap_server
        self.email_address = email_address
        self.password = password
        self.smtp_conn = None
        self.imap_conn = None

    def smtp_connect(self):
        try:
            self.smtp_conn = smtplib.SMTP(self.smtp_server, 587)
            self.smtp_conn.starttls()
            self.smtp_conn.login(self.email_address, self.password)
            print(f"{self.service}: SMTP connection established.")
        except Exception as e:
            return print(f"{self.service}: Error in SMTP connection - {e}")

    def imap_connect(self):
        try:
            self.imap_conn = imaplib.IMAP4_SSL(self.imap_server)
            self.imap_conn.login(self.email_address, self.password)
            print(f"{self.service}: IMAP connection established.")
        except Exception as e:
            return print(f"{self.service}: Error in IMAP connection - {e}")

    def disconnect_smtp(self):
        if self.smtp_conn:
            self.smtp_conn.quit()
            print(f"{self.service}: SMTP connection closed.")

    def disconnect_imap(self):
        if self.imap_conn:
            self.imap_conn.logout()
            print(f"{self.service}: IMAP connection closed.")

    def send_email(self, recipient, subject, body, attachment_path=None):
        try:
            EmailAutomation.imap_connect(self)
            EmailAutomation.smtp_connect(self)
            msg = MIMEMultipart()
            msg["From"] = self.email_address
            msg["To"] = recipient
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain"))
            if attachment_path:
                if not os.path.isfile(attachment_path):
                    print(f"Attachment file does not exist: {attachment_path}")
                elif os.path.getsize(attachment_path) > 25 * 1024 * 1024:  # 25 MB limit
                    body += f"\n\nThe file is too large to attach directly. Please find it in the shared drive."
                    # Here we can add a logic to upload the file to cloud and attach the link
                else:
                    with open(attachment_path, 'rb') as attachment:
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header('Content-Disposition', f'attachment; filename={os.path.basename(attachment_path)}')
                        msg.attach(part)

            self.smtp_conn.sendmail(self.email_address, recipient, msg.as_string())
            print(f"{self.service}: Email sent to {recipient}")
            EmailAutomation.disconnect_imap(self)
            EmailAutomation.disconnect_smtp(self)
            sys.exit(0)
        except Exception as e:
            print(f"{self.service}: Failed to send email - {e}")

    def schedule_email(self,recipient,subject,body,schedule_time,attachment_Path=None):
        def job():
            try:
                EmailAutomation.send_email(self, recipient, subject, body, attachment_Path)
            except Exception as e:
                print(f"{self.service}: Failed to send email - {e}")
            finally:

                schedule.cancel_job(job_instance)  # Cancel the job after sending
                sys.exit(0)
        job_instance = schedule.every().day.at(schedule_time).do(job)

        print(f"Scheduled email to {recipient} at {schedule_time}")

        while True:
            schedule.run_pending()
            time.sleep(1)

    def searchMail(self,searchNumber,senderMail=None,subject=None,attachment_name=None,from_date=None,till_date=None):
        try:
            EmailAutomation.imap_connect(self)
            EmailAutomation.smtp_connect(self)
            self.imap_conn.select("inbox")  # Select the inbox to search
            # Construct the search query
            criteria = []
            if senderMail:
                criteria.append(f'FROM "{senderMail}"')
            if subject:
                criteria.append(f'SUBJECT "{subject}"')
            if from_date:
                # Format the date as IMAP expects 'DD-Mon-YYYY'
                from_date = datetime.datetime.strptime(from_date, "%d-%m-%Y")
                from_date_str = from_date.strftime("%d-%b-%Y")
                criteria.append(f'SINCE "{from_date_str}"')
            if till_date:
                till_date = datetime.datetime.strptime(till_date, "%d-%m-%Y")
                till_date_str = till_date.strftime("%d-%b-%Y")
                criteria.append(f'BEFORE "{till_date_str}"')

            search_criteria = " ".join(criteria)

            # Perform the search
            status, messages = self.imap_conn.search(None, search_criteria)
            if status != "OK":
                print("No messages found with the given criteria.")
                return None

            email_ids = messages[0].split()
            last_n_search_match = email_ids[-searchNumber:]

            # Create a document to store the results
            doc = Document()
            doc.add_heading("Last N Emails", level=1)

            for i in last_n_search_match:
                status, msg_data = self.imap_conn.fetch(i, "(RFC822)")
                msg = email.message_from_bytes(msg_data[0][1])
                subject = msg["Subject"]
                from_ = msg["From"]
                date_ = msg["Date"]
                body_ = msg["Body"]

                # Check for attachment if the user provided an attachment name
                if attachment_name:
                    attachment_found = False
                    for part in msg.walk():
                        # If the part is an attachment and its filename matches
                        if part.get_content_disposition() == "attachment":
                            filename, _ = decode_header(part.get_filename())[0]
                            if isinstance(filename, bytes):
                                filename = filename.decode()
                            if attachment_name in filename:
                                attachment_found = True
                                break
                    if not attachment_found:
                        continue  # Skip this email if no matching attachment is found

                # Add email details to the document
                doc.add_paragraph(f"Subject: {subject}\nFrom: {from_}\nDate: {date_}\n Body: {body_}\n\n")
                doc.add_paragraph()

            # Save the document with the results
            doc.save("result.docx")
            print("The searched emails are saved in 'result.docx'.")
            EmailAutomation.disconnect_imap(self)
            EmailAutomation.disconnect_smtp(self)
        except Exception as e:
            print(f"Error retrieving emails: {e}")

    def getAttachments(self, attachment_count):
        try:
            self.imap_conn.select("inbox")
            status, messages = self.imap_conn.search(None, "ALL")
            email_ids = messages[0].split()
            last_n_attachments = email_ids[-attachment_count:]

            attachments_dir = os.path.join(os.getcwd(), "attachments")
            os.makedirs(attachments_dir, exist_ok=True)

            attachments = []
            for email_id in last_n_attachments:
                status, msg_data = self.imap_conn.fetch(email_id, "(RFC822)")
                msg = email.message_from_bytes(msg_data[0][1])

                # Loop through the email parts to find attachments
                for part in msg.walk():
                    if part.get_content_maintype() == "multipart":
                        continue
                    if part.get("Content-Disposition") is None:
                        continue

                    filename = part.get_filename()
                    if filename:
                        attachments.append((filename, part.get_payload(decode=True)))

            # Download attachments
            for filename, content in attachments:
                file_path = os.path.join(
                    attachments_dir, filename
                )  # Save in the attachments folder
                with open(file_path, "wb") as f:
                    f.write(content)
                    print(f"Downloaded: {file_path}")

            if not attachments:
                print("No attachments found.")
        except Exception as e:
            print(f"Error retrieving attachments: {e}")

def main():
    automation = EmailAutomation(**gmail_config)
    print("Welcome to the Email Automation Tool!")
    print("Choose an option:\n1. Send email \n2. Schedule email"+
            "\n3. Search Email\n4. Download attachment")
    choice = input("Enter your choice: ")
    match choice:
        case "1":
            print("Send Email")
            to_email = input("Enter recipient's email: ")
            subject = input("Enter subject: ")
            body = input("Enter body of the email: ")
            attachment_bool = input("do you have attachment? Y or N: ")
            if attachment_bool == "Y":
                attachment_path = input("Enter attachment path: ")
            elif attachment_bool == "N":
                attachment_path = None
            automation.send_email(to_email, subject, body, attachment_path)
        case "2":
            print("Schedule Email Menu:")
            to_email = input("Enter recipient's email: ")
            subject = input("Enter subject: ")
            body = input("Enter body of the email: ")
            # Get date and time for scheduling
            date_str = input("Enter the date to send the email (YYYY-MM-DD): ")
            time_str = input("Enter the time to send the email (HH:MM in 24-hour format): ")

            schedule_time = f"{date_str} {time_str}"
            schedule_time = datetime.datetime.strptime(schedule_time, "%Y-%m-%d %H:%M")
            # Calculate the time remaining until the scheduled time
            current_time = datetime.datetime.now()
            if schedule_time < current_time:
                print("The scheduled time must be in the future. Please enter a valid date and time.")
                return
            # Schedule the email
            attachment_bool = input("do you have attachment? Y or N: ")
            if attachment_bool == "Y":
                attachment_path = input("Enter attachment path: ")
            elif attachment_bool == "N":
                attachment_path = None
            automation.schedule_email(to_email, subject, body, time_str, attachment_path)

        case "3":
            print("Search Email")
            nums = input("Number of most recent mails you want to search from : ")
            sender = input("Enter senders mail to search : ")
            sub = input("Enter subject to search : ")
            attachment_name = input("Enter attachment name to search from : ")
            from_date = input("Enter date to search from (DD-MM-YYY) : ")
            till_date = input("Enter date to search till (DD-MM-YYY) : ")
            automation.searchMail(int(nums),sender,sub,attachment_name,from_date,till_date)

        case "4":
            print("Search Attachment Menu:")
        case _:
            print("Enter valid input")

if __name__ == "__main__":
    main()

"""def searchBySubject(self, subject, searchNumber):
        try:
            self.imap_conn.select("inbox")
            status, messages = self.imap_conn.search(
                None, "SUBJECT", '"' + subject + '"'
            )
            email_ids = messages[0].split()
            last_n_search_match = email_ids[-searchNumber:]
            doc = Document()
            doc._body.clear_content()
            doc.add_heading("Last N Emails", level=1)

            for i in last_n_search_match:
                status, msg_data = self.imap_conn.fetch(i, "(RFC822)")
                msg = email.message_from_bytes(msg_data[0][1])
                subject = msg["Subject"]
                from_ = msg["From"]
                date_ = msg["Date"]

                doc.add_paragraph(
                    f"Subject : {subject}\nfrom : {from_}\ndate : {date_}\n"
                )
                doc.add_paragraph()

            doc.save("result.docx")
            print("the searched mails are returned in mail search results")
        except Exception as e:
            print(f"Error retrieving emails: {e}")
            return None
            
            
               
    def getAttachmentsBySender(self, senderMail, attachment_count):
        try:
            self.imap_conn.select("inbox")
            messages = self.imap_conn.search(None, f'FROM "{senderMail}"')
            email_ids = messages[0].split()
            last_n_attachments = email_ids[-attachment_count:]

            attachments_dir = os.path.join(os.getcwd(), "attachments")
            os.makedirs(attachments_dir, exist_ok=True)

            attachments = []
            for email_id in last_n_attachments:
                msg_data = self.imap_conn.fetch(email_id, "(RFC822)")
                msg = email.message_from_bytes(msg_data[0][1])

                # Loop through the email parts to find attachments
                for part in msg.walk():
                    if part.get_content_maintype() == "multipart":
                        continue
                    if part.get("Content-Disposition") is None:
                        continue

                    filename = part.get_filename()
                    if filename:
                        attachments.append((filename, part.get_payload(decode=True)))

            # Download attachments
            for filename, content in attachments:
                file_path = os.path.join(
                    attachments_dir, filename
                )  # Save in the attachments folder
                with open(file_path, "wb") as f:
                    f.write(content)
                    print(f"Downloaded: {file_path}")

            if not attachments:
                print("No attachments found.")
        except Exception as e:
            print(f"Error retrieving attachments: {e}")

def searchMail(self, senderMail, searchNumber):
        try:
            EmailAutomation.imap_connect()
            EmailAutomation.smtp_connect()
            self.imap_conn.select("inbox")
            status, messages = self.imap_conn.search(None, f'FROM "{senderMail}"')
            email_ids = messages[0].split()
            last_n_search_match = email_ids[-searchNumber:]
            doc = Document()
            doc._body.clear_content()
            doc.add_heading("Last N Emails", level=1)

            for i in last_n_search_match:
                status, msg_data = self.imap_conn.fetch(i, "(RFC822)")
                msg = email.message_from_bytes(msg_data[0][1])
                subject = msg["Subject"]
                from_ = msg["From"]
                date_ = msg["Date"]

                doc.add_paragraph(
                    f"Subject : {subject}\nfrom : {from_}\ndate : {date_}\n"
                )
                doc.add_paragraph()

            doc.save("result.docx")
            print("the searched mails are returned in mail search results")
        except Exception as e:
            print(f"Error retrieving emails: {e}")
            return None
"""
